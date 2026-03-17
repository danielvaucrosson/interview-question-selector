"""
Docx Generation Engine for Interview Question Sets.

Ports the interview-prep-template.js formatting to Python using python-docx.
Generates a fully formatted .docx interview document with cover page,
question sections, and overall assessment.
"""

from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Style constants
# ---------------------------------------------------------------------------
DARK_BLUE = "1F3864"
MED_BLUE = "2E75B6"
LIGHT_BLUE = "D6E4F0"
LIGHT_GRAY = "F2F2F2"

DIFF_LABELS = {0: "Intro", 1: "Foundational", 2: "Applied", 3: "Architectural"}
DIFF_COLORS = {
    0: "2E75B6",
    1: "70AD47",
    2: "ED7D31",
    3: "C00000",
}

DEFAULT_ASSESSMENT_DIMENSIONS = [
    "Technical Depth",
    "Data Modeling & Architecture Understanding",
    "Platform & Tooling Experience",
    "Client Communication & Storytelling",
    "Problem-Solving & Consulting Judgment",
    "Pre-Sales / POC Capability",
]

DEFAULT_EVAL_DIMENSIONS = [
    {
        "label": "Tech Skills",
        "prompt": "What are the main technical strengths you observed in the candidate during the interview?",
        "has_rating": True,
    },
    {
        "label": "Consulting Skills",
        "prompt": (
            "How would you rate the candidate\u2019s consulting skills, particularly in their"
            " capacity to comprehend client requirements and deliver impactful solutions? Why?"
        ),
        "has_rating": True,
    },
    {
        "label": "Seniority",
        "prompt": "Does the candidate possess the appropriate seniority for the role, in your assessment?",
        "has_rating": True,
    },
    {
        "label": "Communication",
        "prompt": "How would you describe the candidate\u2019s communication style during the interview?",
        "has_rating": True,
    },
    {
        "label": "Cultural Fit",
        "prompt": (
            "In your opinion, does the candidate demonstrate a good cultural fit with the team"
            " and the organization? Why?"
        ),
        "has_rating": True,
    },
    {
        "label": "Weaknesses / Areas for Improvement",
        "prompt": "Please share your insights into the candidate\u2019s identified weaknesses or areas for improvement.",
        "has_rating": False,
    },
]


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, color: str):
    """Apply background shading to a table cell."""
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    )


def _set_cell_borders(cell, color: str = "CCCCCC", size: int = 4):
    """Apply uniform single borders to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(
        parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="{size}" w:color="{color}"/>'
            f'  <w:bottom w:val="single" w:sz="{size}" w:color="{color}"/>'
            f'  <w:left w:val="single" w:sz="{size}" w:color="{color}"/>'
            f'  <w:right w:val="single" w:sz="{size}" w:color="{color}"/>'
            f'</w:tcBorders>'
        )
    )


def _set_cell_no_borders(cell):
    """Remove all borders from a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(
        parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="none" w:sz="0" w:color="auto"/>'
            f'  <w:bottom w:val="none" w:sz="0" w:color="auto"/>'
            f'  <w:left w:val="none" w:sz="0" w:color="auto"/>'
            f'  <w:right w:val="none" w:sz="0" w:color="auto"/>'
            f'</w:tcBorders>'
        )
    )


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    """Set cell margins in twips."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tc_pr.append(tc_mar)


def _set_row_height(row, val: int, rule: str = "atLeast"):
    """Set table row height. val is in twips."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trPr.append(
        parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{val}" w:hRule="{rule}"/>')
    )


def _add_run(paragraph, text, size_pt=10, bold=False, italic=False, color=None, font="Arial"):
    """Add a styled run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def _add_paragraph_shading(paragraph, color: str):
    """Apply background shading to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'))


def _add_paragraph_bottom_border(paragraph, color: str = MED_BLUE, size: int = 12):
    """Add a bottom border to a paragraph (used as separator)."""
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(
        parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="{size}" w:color="{color}" w:space="1"/>'
            f'</w:pBdr>'
        )
    )


# ---------------------------------------------------------------------------
# Table construction helpers
# ---------------------------------------------------------------------------

def _make_info_table(doc, metadata):
    """Create the candidate info table on the cover page (5 rows, 2 cols)."""
    rows_data = [
        ("Candidate:", metadata.get("candidate_name", "")),
        ("Date:", metadata.get("interview_date", "")),
        ("Interviewer:", metadata.get("interviewer", "")),
        ("Questions:", str(metadata.get("question_count", ""))),
        ("Est. Duration:", metadata.get("duration", "")),
    ]
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    for row_idx, (label, value) in enumerate(rows_data):
        row = table.rows[row_idx]
        # Label cell
        lc = row.cells[0]
        lc.width = Emu(2400 * 914)
        lc.paragraphs[0].clear()
        _add_run(lc.paragraphs[0], label, size_pt=10, bold=True)
        _set_cell_shading(lc, LIGHT_GRAY)
        _set_cell_borders(lc)
        _set_cell_margins(lc)

        # Value cell
        vc = row.cells[1]
        vc.width = Emu(3600 * 914)
        vc.paragraphs[0].clear()
        _add_run(vc.paragraphs[0], value, size_pt=10)
        _set_cell_borders(vc)
        _set_cell_margins(vc)

    return table


def _make_meta_table(doc, q):
    """Create the 5-column metadata table for a question."""
    diff_val = q.get("difficulty", 0)
    diff_label = DIFF_LABELS.get(diff_val, "Unknown")
    diff_color = DIFF_COLORS.get(diff_val, "000000")

    table = doc.add_table(rows=2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    headers = ["ID", "Domain", "Technology", "Difficulty", "Type"]
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].clear()
        _add_run(cell.paragraphs[0], h, size_pt=10, bold=True, color="FFFFFF")
        _set_cell_shading(cell, DARK_BLUE)
        _set_cell_borders(cell)
        _set_cell_margins(cell)

    # Data row
    data_cells = [
        (q.get("id", ""), 10, True, None),
        (q.get("subdomain", ""), 9, False, None),
        (q.get("technology", "General"), 9, False, None),
        (f"{diff_val} - {diff_label}", 9, True, diff_color),
        (q.get("question_type", ""), 9, False, None),
    ]
    for i, (text, size, bold, color) in enumerate(data_cells):
        cell = table.rows[1].cells[i]
        cell.paragraphs[0].clear()
        run = _add_run(cell.paragraphs[0], text, size_pt=size, bold=bold, color=color)
        if i == 4:  # Type column is italic
            run.italic = True
        _set_cell_borders(cell)
        _set_cell_margins(cell)

    return table


def _make_notes_box(doc, height_twips=1200, border_color="999999"):
    """Create a 1-row, 1-column bordered notes box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    cell.paragraphs[0].clear()
    _add_run(cell.paragraphs[0], "", size_pt=10)
    _set_cell_borders(cell, color=border_color, size=4)
    _set_cell_shading(cell, "FAFAFA")
    _set_cell_margins(cell)
    _set_row_height(table.rows[0], height_twips)
    return table


def _make_rating_row(doc):
    """Create the rating row: label + options on left, empty input box on right."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Left cell — no borders
    left = table.rows[0].cells[0]
    left.paragraphs[0].clear()
    _add_run(left.paragraphs[0], "Rating:  ", size_pt=10, bold=True)
    _add_run(
        left.paragraphs[0],
        "1 (Weak)  2 (Below)  3 (Meets)  4 (Strong)  5 (Exceptional)",
        size_pt=9,
    )
    _set_cell_no_borders(left)
    _set_cell_margins(left, top=60, bottom=60, left=0, right=120)

    # Right cell — bordered input box
    right = table.rows[0].cells[1]
    right.paragraphs[0].clear()
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(right.paragraphs[0], "", size_pt=10)
    _set_cell_borders(right, color="999999")
    _set_cell_shading(right, "FAFAFA")
    _set_cell_margins(right, top=60, bottom=60, left=120, right=120)

    return table


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _build_cover_page(doc, metadata, sections):
    """Build the cover/title page."""
    # Spacer
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(150)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(10)
    _add_run(title_p, "Interview Question Set", size_pt=24, bold=True, color=DARK_BLUE)

    # Job title subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(6)
    _add_run(sub_p, metadata.get("job_title", ""), size_pt=14, color=MED_BLUE)

    # Company — Seniority
    comp_p = doc.add_paragraph()
    comp_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    comp_p.paragraph_format.space_after = Pt(30)
    company = metadata.get("company", "")
    seniority = metadata.get("seniority", "")
    _add_run(comp_p, f"{company} \u2014 {seniority}", size_pt=12, color="444444")

    # Info table
    _make_info_table(doc, metadata)

    # Spacer
    sp2 = doc.add_paragraph()
    sp2.paragraph_format.space_before = Pt(30)

    # Difficulty distribution
    all_questions = [q for s in sections for q in s.get("questions", [])]
    diff_counts = {0: 0, 1: 0, 2: 0, 3: 0}
    for q in all_questions:
        d = q.get("difficulty", 0)
        diff_counts[d] = diff_counts.get(d, 0) + 1

    dist_p = doc.add_paragraph()
    dist_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dist_p.paragraph_format.space_after = Pt(4)
    _add_run(dist_p, "Difficulty Distribution:  ", size_pt=10)

    parts = []
    if diff_counts[0]:
        parts.append((f"{diff_counts[0]} Intro", "2E75B6"))
    if diff_counts[1]:
        parts.append((f"{diff_counts[1]} Foundational", "70AD47"))
    if diff_counts[2]:
        parts.append((f"{diff_counts[2]} Applied", "ED7D31"))
    if diff_counts[3]:
        parts.append((f"{diff_counts[3]} Architectural", "C00000"))

    for idx, (text, color) in enumerate(parts):
        _add_run(dist_p, text, size_pt=10, bold=True, color=color)
        if idx < len(parts) - 1:
            _add_run(dist_p, "  \u2022  ", size_pt=10, color="CCCCCC")


def _build_question_block(doc, q):
    """Build all elements for a single question."""
    # Metadata table
    _make_meta_table(doc, q)

    # Small spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(6)

    # Question text
    q_p = doc.add_paragraph()
    q_p.paragraph_format.space_before = Pt(4)
    q_p.paragraph_format.space_after = Pt(4)
    _add_run(q_p, "Q: ", size_pt=11, bold=True, color=MED_BLUE)
    _add_run(q_p, q.get("question", ""), size_pt=11)

    # Ideal Answer Components header
    ans_header = doc.add_paragraph()
    ans_header.paragraph_format.space_before = Pt(6)
    ans_header.paragraph_format.space_after = Pt(2)
    _add_paragraph_shading(ans_header, LIGHT_BLUE)
    _add_run(ans_header, "  Ideal Answer Components:", size_pt=10, bold=True, color=DARK_BLUE)

    # Answer bullets
    answer_text = q.get("answer", "")
    bullets = [b.strip() for b in answer_text.split("\n") if b.strip()]
    for bullet in bullets:
        bp = doc.add_paragraph()
        bp.paragraph_format.space_before = Pt(2)
        bp.paragraph_format.space_after = Pt(2)
        bp.paragraph_format.left_indent = Inches(0.25)
        _add_run(bp, bullet, size_pt=9.5)

    # Notes label
    notes_label = doc.add_paragraph()
    notes_label.paragraph_format.space_before = Pt(6)
    notes_label.paragraph_format.space_after = Pt(2)
    _add_run(notes_label, "Notes:", size_pt=10, bold=True)

    # Notes box
    _make_notes_box(doc, height_twips=1200)

    # Rating row
    _make_rating_row(doc)

    # Small spacer
    sp2 = doc.add_paragraph()
    sp2.paragraph_format.space_after = Pt(6)

    # Separator line
    sep = doc.add_paragraph()
    sep.paragraph_format.space_after = Pt(10)
    _add_paragraph_bottom_border(sep, color=MED_BLUE, size=12)


def _build_question_sections(doc, sections):
    """Build all question sections with headers."""
    # Page break before questions
    doc.add_page_break()

    for section in sections:
        questions = section.get("questions", [])
        if not questions:
            continue

        # Section header with dark blue background
        header_p = doc.add_paragraph()
        header_p.paragraph_format.space_before = Pt(10)
        header_p.paragraph_format.space_after = Pt(4)
        _add_paragraph_shading(header_p, DARK_BLUE)
        _add_run(
            header_p,
            f"  {section.get('title', '')}",
            size_pt=12,
            bold=True,
            color="FFFFFF",
        )

        # Section subtitle
        sub_p = doc.add_paragraph()
        sub_p.paragraph_format.space_after = Pt(10)
        _add_run(sub_p, section.get("subtitle", ""), size_pt=9.5, italic=True)

        # Questions
        for q in questions:
            _build_question_block(doc, q)


def _build_assessment_page(doc, dimensions=None):
    """Build the overall assessment page."""
    if dimensions is None:
        dimensions = DEFAULT_ASSESSMENT_DIMENSIONS

    doc.add_page_break()

    # Header
    header_p = doc.add_paragraph()
    header_p.paragraph_format.space_after = Pt(10)
    _add_paragraph_shading(header_p, DARK_BLUE)
    _add_run(header_p, "  OVERALL ASSESSMENT", size_pt=12, bold=True, color="FFFFFF")

    # Assessment dimensions
    for label in dimensions:
        # 3-column table: label | "1-5:" | input box
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Label cell
        lc = table.rows[0].cells[0]
        lc.paragraphs[0].clear()
        _add_run(lc.paragraphs[0], label, size_pt=9.5, bold=True)
        # Borders on top/bottom/left, no right
        tc_pr = lc._tc.get_or_add_tcPr()
        tc_pr.append(
            parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:top w:val="single" w:sz="4" w:color="999999"/>'
                f'  <w:bottom w:val="single" w:sz="4" w:color="999999"/>'
                f'  <w:left w:val="single" w:sz="4" w:color="999999"/>'
                f'  <w:right w:val="none" w:sz="0" w:color="auto"/>'
                f'</w:tcBorders>'
            )
        )
        _set_cell_margins(lc, top=60, bottom=60, left=120, right=60)

        # "1-5:" cell
        mc = table.rows[0].cells[1]
        mc.paragraphs[0].clear()
        mc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_run(mc.paragraphs[0], "1\u20135:", size_pt=9)
        tc_pr2 = mc._tc.get_or_add_tcPr()
        tc_pr2.append(
            parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:top w:val="single" w:sz="4" w:color="999999"/>'
                f'  <w:bottom w:val="single" w:sz="4" w:color="999999"/>'
                f'  <w:left w:val="none" w:sz="0" w:color="auto"/>'
                f'  <w:right w:val="none" w:sz="0" w:color="auto"/>'
                f'</w:tcBorders>'
            )
        )
        _set_cell_margins(mc, top=60, bottom=60, left=60, right=60)

        # Input box cell
        rc = table.rows[0].cells[2]
        rc.paragraphs[0].clear()
        rc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(rc.paragraphs[0], "", size_pt=10)
        _set_cell_borders(rc, color="999999")
        _set_cell_shading(rc, "FAFAFA")
        _set_cell_margins(rc, top=60, bottom=60, left=120, right=120)

        # Notes box (smaller)
        _make_notes_box(doc, height_twips=600)

        # Small spacer
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(4)

    # Overall Recommendation
    rec_label = doc.add_paragraph()
    rec_label.paragraph_format.space_before = Pt(15)
    rec_label.paragraph_format.space_after = Pt(3)
    _add_run(rec_label, "Overall Recommendation:", size_pt=11, bold=True)

    rec_table = doc.add_table(rows=1, cols=2)
    rec_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Options cell (no borders)
    opt_cell = rec_table.rows[0].cells[0]
    opt_cell.paragraphs[0].clear()
    _add_run(opt_cell.paragraphs[0], "Strong Hire  /  Hire  /  Maybe  /  No Hire", size_pt=10)
    _set_cell_no_borders(opt_cell)
    _set_cell_margins(opt_cell, top=60, bottom=60, left=0, right=120)

    # Input box cell
    inp_cell = rec_table.rows[0].cells[1]
    inp_cell.paragraphs[0].clear()
    inp_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(inp_cell.paragraphs[0], "", size_pt=10)
    _set_cell_borders(inp_cell, color="999999")
    _set_cell_shading(inp_cell, "FAFAFA")
    _set_cell_margins(inp_cell, top=60, bottom=60, left=120, right=120)

    # Summary Notes
    notes_label = doc.add_paragraph()
    notes_label.paragraph_format.space_before = Pt(10)
    notes_label.paragraph_format.space_after = Pt(3)
    _add_run(notes_label, "Summary Notes:", size_pt=10, bold=True)

    # Large notes box (~1.7" = ~2448 twips)
    _make_notes_box(doc, height_twips=2400)


def _build_eval_dimension(doc, dim):
    """Build a single evaluation dimension block (header bar, prompt, optional rating, notes)."""
    # Light-blue header bar with dimension label
    hdr = doc.add_paragraph()
    hdr.paragraph_format.space_before = Pt(10)
    hdr.paragraph_format.space_after = Pt(4)
    _add_paragraph_shading(hdr, LIGHT_BLUE)
    _add_run(hdr, f"  {dim['label']}", size_pt=10, bold=True, color=DARK_BLUE)

    # Prompt text
    prompt_p = doc.add_paragraph()
    prompt_p.paragraph_format.space_before = Pt(2)
    prompt_p.paragraph_format.space_after = Pt(4)
    _add_run(prompt_p, dim["prompt"], size_pt=10)

    # Rating row (if applicable)
    if dim.get("has_rating", False):
        rating_lbl = doc.add_paragraph()
        rating_lbl.paragraph_format.space_before = Pt(2)
        rating_lbl.paragraph_format.space_after = Pt(2)
        _make_rating_row(doc)

    # Notes label + notes box
    notes_lbl = doc.add_paragraph()
    notes_lbl.paragraph_format.space_before = Pt(4)
    notes_lbl.paragraph_format.space_after = Pt(2)
    _add_run(notes_lbl, "Notes:", size_pt=10, bold=True)

    _make_notes_box(doc, height_twips=700)


def _build_ats_evaluation_page(doc, eval_dimensions=None):
    """Build the ATS Evaluation page (Page 1 of the evaluation section)."""
    if eval_dimensions is None:
        eval_dimensions = DEFAULT_EVAL_DIMENSIONS

    doc.add_page_break()

    # Dark blue header bar
    header_p = doc.add_paragraph()
    header_p.paragraph_format.space_after = Pt(4)
    _add_paragraph_shading(header_p, DARK_BLUE)
    _add_run(
        header_p,
        "  ATS EVALUATION \u2014 Technical Interview | General Template",
        size_pt=12,
        bold=True,
        color="FFFFFF",
    )

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(10)
    _add_run(
        sub_p,
        "Complete this section after the interview and transfer responses to your ATS.",
        size_pt=9.5,
        italic=True,
    )

    # Evaluation dimensions
    for dim in eval_dimensions:
        _build_eval_dimension(doc, dim)


def _build_internal_only_page(doc):
    """Build the Internal Only page (Page 2 of the evaluation section)."""
    doc.add_page_break()

    # Red header bar
    header_p = doc.add_paragraph()
    header_p.paragraph_format.space_after = Pt(10)
    _add_paragraph_shading(header_p, "C00000")
    _add_run(header_p, "  INTERNAL ONLY", size_pt=12, bold=True, color="FFFFFF")

    # --- Seniority Classification ---
    sen_hdr = doc.add_paragraph()
    sen_hdr.paragraph_format.space_before = Pt(10)
    sen_hdr.paragraph_format.space_after = Pt(4)
    _add_paragraph_shading(sen_hdr, LIGHT_BLUE)
    _add_run(sen_hdr, "  [Internal] Seniority Classification", size_pt=10, bold=True, color=DARK_BLUE)

    sen_prompt = doc.add_paragraph()
    sen_prompt.paragraph_format.space_before = Pt(2)
    sen_prompt.paragraph_format.space_after = Pt(4)
    _add_run(
        sen_prompt,
        "Based on the candidate\u2019s experience and responses, how would you classify their seniority level?",
        size_pt=10,
    )

    # Options row
    sen_table = doc.add_table(rows=1, cols=2)
    sen_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    opt_cell = sen_table.rows[0].cells[0]
    opt_cell.paragraphs[0].clear()
    _add_run(
        opt_cell.paragraphs[0],
        "Junior  /  Mid  /  Semi Senior  /  Senior  /  Expert/Technical Lead  /  Manager/Director",
        size_pt=10,
    )
    _set_cell_no_borders(opt_cell)
    _set_cell_margins(opt_cell, top=60, bottom=60, left=0, right=120)

    inp_cell = sen_table.rows[0].cells[1]
    inp_cell.paragraphs[0].clear()
    inp_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(inp_cell.paragraphs[0], "", size_pt=10)
    _set_cell_borders(inp_cell, color="999999")
    _set_cell_shading(inp_cell, "FAFAFA")
    _set_cell_margins(inp_cell, top=60, bottom=60, left=120, right=120)

    # --- Recommendation ---
    rec_hdr = doc.add_paragraph()
    rec_hdr.paragraph_format.space_before = Pt(14)
    rec_hdr.paragraph_format.space_after = Pt(4)
    _add_paragraph_shading(rec_hdr, LIGHT_BLUE)
    _add_run(rec_hdr, "  [Internal] Recommendation", size_pt=10, bold=True, color=DARK_BLUE)

    rec_prompt = doc.add_paragraph()
    rec_prompt.paragraph_format.space_before = Pt(2)
    rec_prompt.paragraph_format.space_after = Pt(4)
    _add_run(
        rec_prompt,
        "Would you recommend moving this candidate forward to the next step in the hiring process?",
        size_pt=10,
    )

    rec_table = doc.add_table(rows=1, cols=2)
    rec_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    rec_opt = rec_table.rows[0].cells[0]
    rec_opt.paragraphs[0].clear()
    _add_run(rec_opt.paragraphs[0], "Yes  /  No", size_pt=10)
    _set_cell_no_borders(rec_opt)
    _set_cell_margins(rec_opt, top=60, bottom=60, left=0, right=120)

    rec_inp = rec_table.rows[0].cells[1]
    rec_inp.paragraphs[0].clear()
    rec_inp.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(rec_inp.paragraphs[0], "", size_pt=10)
    _set_cell_borders(rec_inp, color="999999")
    _set_cell_shading(rec_inp, "FAFAFA")
    _set_cell_margins(rec_inp, top=60, bottom=60, left=120, right=120)

    # --- Feedback and Comments ---
    fb_hdr = doc.add_paragraph()
    fb_hdr.paragraph_format.space_before = Pt(14)
    fb_hdr.paragraph_format.space_after = Pt(4)
    _add_paragraph_shading(fb_hdr, LIGHT_BLUE)
    _add_run(fb_hdr, "  [Internal] Feedback and Comments", size_pt=10, bold=True, color=DARK_BLUE)

    fb_prompt = doc.add_paragraph()
    fb_prompt.paragraph_format.space_before = Pt(2)
    fb_prompt.paragraph_format.space_after = Pt(4)
    _add_run(fb_prompt, "Any additional comments or observations from the interview.", size_pt=10)

    # Large notes box
    _make_notes_box(doc, height_twips=2400)


def _setup_header_footer(doc, metadata):
    """Configure document header and footer."""
    section = doc.sections[0]

    # Header
    header = section.header
    header.is_linked_to_previous = False
    header_p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_p.clear()
    company = metadata.get("company", "")
    candidate = metadata.get("candidate_name", "")
    _add_run(header_p, f"{company} \u2014 Interview: {candidate}", size_pt=8, color="999999")
    # Bottom border on header paragraph
    pPr = header_p._p.get_or_add_pPr()
    pPr.append(
        parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="8" w:color="{MED_BLUE}" w:space="4"/>'
            f'</w:pBdr>'
        )
    )

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_p.clear()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # "Page " text
    run1 = footer_p.add_run("Page ")
    run1.font.name = "Arial"
    run1.font.size = Pt(8)
    run1.font.color.rgb = RGBColor.from_string("999999")

    # Page number field
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    fld_instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')

    run2 = footer_p.add_run()
    run2.font.name = "Arial"
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor.from_string("999999")
    run2._r.append(fld_char_begin)

    run3 = footer_p.add_run()
    run3.font.name = "Arial"
    run3.font.size = Pt(8)
    run3.font.color.rgb = RGBColor.from_string("999999")
    run3._r.append(fld_instr)

    run4 = footer_p.add_run()
    run4.font.name = "Arial"
    run4.font.size = Pt(8)
    run4.font.color.rgb = RGBColor.from_string("999999")
    run4._r.append(fld_char_end)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def generate_interview_docx(
    metadata: dict, sections: list[dict], eval_dimensions: list[dict] | None = None
) -> BytesIO:
    """
    Generate a formatted interview document.

    Args:
        metadata: dict with keys: candidate_name, interview_date, interviewer,
                  job_title, company, seniority, duration, question_count
        sections: list of dicts, each with keys: title, subtitle, questions
                  Each question dict has: id, question, answer, subdomain,
                  technology, difficulty, question_type
        eval_dimensions: optional list of dicts, each with keys: label, prompt,
                  has_rating.  Defaults to DEFAULT_EVAL_DIMENSIONS.

    Returns:
        BytesIO buffer containing the .docx file
    """
    doc = Document()

    # Page setup — Letter 8.5" x 11", 1" margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Default font style
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Build document sections
    _build_cover_page(doc, metadata, sections)
    _build_question_sections(doc, sections)
    _build_assessment_page(doc)
    _build_ats_evaluation_page(doc, eval_dimensions)
    _build_internal_only_page(doc)

    # Header & footer
    _setup_header_footer(doc, metadata)

    # Save to BytesIO
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
